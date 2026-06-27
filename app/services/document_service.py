from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Any

from fastapi import UploadFile


@dataclass
class ExtractedDocument:
    filename: str
    extension: str
    text: str
    char_count: int
    chunk_count: int


class DocumentService:
    """
    Membaca dokumen untuk fitur Document Intelligence.

    Supported:
    - .txt / .md
    - .docx via python-docx
    - .pdf via pypdf

    Target utama saat ini:
    - Laporan Preventive Maintenance / Health Check CrowdStrike Falcon.
    """

    def __init__(
        self,
        max_file_size_mb: int = 25,
        chunk_size: int = 12000,
        chunk_overlap: int = 800,
    ) -> None:
        self.max_file_size_mb = max_file_size_mb
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    async def extract_from_upload(self, file: UploadFile) -> ExtractedDocument:
        raw = await file.read()
        max_bytes = self.max_file_size_mb * 1024 * 1024

        if len(raw) > max_bytes:
            raise ValueError(f"Ukuran file terlalu besar. Maksimal {self.max_file_size_mb} MB.")

        filename = file.filename or "uploaded_document"
        extension = self._get_extension(filename)

        if extension in {".txt", ".md"}:
            text = self._extract_text(raw)
        elif extension == ".docx":
            text = self._extract_docx(raw)
        elif extension == ".pdf":
            text = self._extract_pdf(raw)
        else:
            raise ValueError(
                f"Format {extension} belum didukung. Gunakan .docx, .pdf, .txt, atau .md."
            )

        text = self.clean_text(text)
        chunks = self.chunk_text(text)

        return ExtractedDocument(
            filename=filename,
            extension=extension,
            text=text,
            char_count=len(text),
            chunk_count=len(chunks),
        )

    def extract_from_text(self, text: str, filename: str = "manual_input.txt") -> ExtractedDocument:
        text = self.clean_text(text)
        chunks = self.chunk_text(text)

        return ExtractedDocument(
            filename=filename,
            extension=".txt",
            text=text,
            char_count=len(text),
            chunk_count=len(chunks),
        )

    def clean_text(self, text: str) -> str:
        text = text.replace("\x00", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        return text.strip()

    def chunk_text(self, text: str) -> list[str]:
        if not text:
            return []

        if len(text) <= self.chunk_size:
            return [text]

        chunks: list[str] = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]

            if end < len(text):
                last_break = chunk.rfind("\n\n")
                if last_break > self.chunk_size * 0.6:
                    chunk = chunk[:last_break]
                    end = start + last_break

            chunks.append(chunk.strip())

            next_start = end - self.chunk_overlap
            if next_start <= start:
                next_start = end
            start = next_start

        return [chunk for chunk in chunks if chunk]

    def compact_document_for_llm(self, text: str, max_chars: int = 60000) -> dict[str, Any]:
        """
        Mencegah context window terlalu besar.
        Jika dokumen panjang, ambil bagian awal + akhir.
        Bagian awal biasanya berisi executive summary dan TOC.
        Bagian akhir biasanya berisi conclusion dan rekomendasi.
        """
        text = self.clean_text(text)

        if len(text) <= max_chars:
            return {"is_truncated": False, "char_count": len(text), "text": text}

        head_size = int(max_chars * 0.65)
        tail_size = max_chars - head_size

        return {
            "is_truncated": True,
            "char_count": len(text),
            "text": (
                text[:head_size]
                + "\n\n...[DOCUMENT TRUNCATED FOR LLM CONTEXT]...\n\n"
                + text[-tail_size:]
            ),
        }

    def _get_extension(self, filename: str) -> str:
        dot = filename.lower().rfind(".")
        return "" if dot == -1 else filename.lower()[dot:]

    def _extract_text(self, raw: bytes) -> str:
        for enc in ("utf-8", "utf-16", "latin-1"):
            try:
                return raw.decode(enc)
            except UnicodeDecodeError:
                pass
        return raw.decode("utf-8", errors="ignore")

    def _extract_docx(self, raw: bytes) -> str:
        try:
            from docx import Document
        except ImportError as error:
            raise RuntimeError("Install dulu: pip install python-docx") from error

        doc = Document(io.BytesIO(raw))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            value = paragraph.text.strip()
            if value:
                parts.append(value)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    def _extract_pdf(self, raw: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as error:
            raise RuntimeError("Install dulu: pip install pypdf") from error

        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
